from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Estilo padrão: Arial 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # --- CAPA ---
    title = doc.add_heading('RELATÓRIO TÉCNICO: TechService OS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Sistema de Gestão de Assistência Técnica e Inventário')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Autor: Lucas Sol Henrique Jacques de Oliveira\n')
    run.bold = True
    info.add_run('Matrícula: 842100049\n')
    info.add_run('Disciplina: Frameworks Web\n')
    info.add_run('Curso: Análise e Desenvolvimento de Sistemas\n')
    info.add_run('Data: 15 de Maio de 2026')

    doc.add_page_break()

    # --- 1. INTRODUÇÃO ---
    doc.add_heading('1. INTRODUÇÃO', level=1)
    p1 = doc.add_paragraph(
        "O presente relatório documenta o desenvolvimento do TechService OS, uma aplicação backend especializada na "
        "gestão de ordens de serviço e controle de inventário para assistências técnicas. O projeto foi concebido para "
        "demonstrar a aplicação prática de padrões arquiteturais avançados e o uso estratégico de frameworks modernos "
        "no ecossistema Node.js."
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2 = doc.add_paragraph(
        "A aplicação utiliza o framework NestJS, seguindo os princípios da Clean Architecture e Domain-Driven Design (DDD). "
        "Esta escolha justifica-se pela necessidade de criar um sistema onde as regras de negócio sejam independentes de "
        "detalhes técnicos (banco de dados, drivers, frameworks), garantindo longevidade e facilidade de manutenção ao código."
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- 2. O PROBLEMA ---
    doc.add_heading('2. O PROBLEMA', level=1)
    doc.add_paragraph(
        "Assistências técnicas frequentemente sofrem com a falta de padronização em seus processos. Os principais desafios identificados foram:"
    )
    doc.add_paragraph("Inconsistência de Status: Dificuldade em rastrear o estágio real de um reparo.", style='List Bullet')
    doc.add_paragraph("Furo de Estoque: Peças são utilizadas sem baixa automática no inventário.", style='List Bullet')
    doc.add_paragraph("Falta de Auditoria: Impossibilidade de rastrear quem alterou dados críticos.", style='List Bullet')
    doc.add_paragraph("Acoplamento Tecnológico: Sistemas que exigem reescrita total ao trocar de tecnologia.", style='List Bullet')

    # --- 3. SOLUÇÃO PROPOSTA ---
    doc.add_heading('3. SOLUÇÃO PROPOSTA', level=1)
    doc.add_paragraph("O TechService OS resolve esses problemas através de uma API REST que centraliza o ciclo de vida da assistência.")
    doc.add_heading('3.1 Requisitos e Funcionalidades', level=2)
    doc.add_paragraph("Gestão de Clientes: CRUD completo com validações de dados únicos (E-mail/CPF).", style='List Bullet')
    doc.add_paragraph("Controle de Peças: Cadastro de componentes com controle de custo e venda.", style='List Bullet')
    doc.add_paragraph("Ordens de Serviço: Vínculo entre cliente, produto e peças necessárias.", style='List Bullet')
    doc.add_paragraph("Reserva de Estoque: Bloqueio de peças para evitar erros de inventário.", style='List Bullet')

    # --- 4. ARQUITETURA DO SISTEMA ---
    doc.add_heading('4. ARQUITETURA DO SISTEMA', level=1)
    doc.add_paragraph("A arquitetura foi estruturada em camadas concêntricas respeitando a Clean Architecture:")
    
    doc.add_heading('4.1 Camadas', level=2)
    doc.add_paragraph("Domain: Contém as Entidades e as Regras de Negócio puras (Máquinas de Estado).", style='List Number')
    doc.add_paragraph("Application: Orquestra o fluxo de dados e define contratos de repositórios.", style='List Number')
    doc.add_paragraph("Infrastructure: Detalhes técnicos como Prisma ORM, PostgreSQL e Redis.", style='List Number')
    doc.add_paragraph("Interface Adapters: Controllers que recebem requisições HTTP via NestJS.", style='List Number')
    
    doc.add_heading('4.2 Decisões Técnicas', level=2)
    doc.add_paragraph("TypeScript: Segurança de tipos em todo o projeto.", style='List Bullet')
    doc.add_paragraph("Injeção de Dependência: Desacoplamento de camadas via container do NestJS.", style='List Bullet')

    # --- 5. FERRAMENTAS UTILIZADAS ---
    doc.add_heading('5. FERRAMENTAS UTILIZADAS', level=1)
    doc.add_paragraph("Stack tecnológica utilizada:")
    doc.add_paragraph("NestJS: Framework principal para modularização.", style='List Bullet')
    doc.add_paragraph("Prisma ORM: Sincronização entre código e SQL.", style='List Bullet')
    doc.add_paragraph("PostgreSQL: Banco de dados relacional robusto.", style='List Bullet')
    doc.add_paragraph("Redis (Conceito A): Camada de Cache para alta performance.", style='List Bullet')
    doc.add_paragraph("Swagger: Documentação automática da API.", style='List Bullet')
    
    doc.add_heading('5.1 Uso Crítico de IA (Relato)', level=2)
    p_ia = doc.add_paragraph(
        "Durante o desenvolvimento, utilizei ferramentas de IA para geração de boilerplate e debugging. "
        "Análise Crítica: Embora a IA acelere o processo, foi necessária uma revisão rigorosa nas sugestões de arquitetura para garantir que a lógica de Domain permanecesse isolada, conforme exigido pela Clean Architecture."
    )
    p_ia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- 6. FUNCIONALIDADES DIFERENCIADAS (CONCEITO A) ---
    doc.add_heading('6. FUNCIONALIDADES DIFERENCIADAS (CONCEITO A)', level=1)
    doc.add_paragraph("Sistema de Auditoria (Audit Logs): Registro automático de mudanças críticas.", style='List Bullet')
    doc.add_paragraph("Cache com Redis: Otimização de consultas de inventário.", style='List Bullet')
    doc.add_paragraph("Tratamento de Exceções Global: Padronização de respostas de erro.", style='List Bullet')

    # --- 7. CONSIDERAÇÕES FINAIS ---
    doc.add_heading('7. CONSIDERAÇÕES FINAIS', level=1)
    p_final = doc.add_paragraph(
        "O desenvolvimento do TechService OS permitiu consolidar conhecimentos de arquitetura avançada. "
        "A separação em camadas provou-se valiosa para a testabilidade e evolução do sistema. "
        "A integração com Redis e auditoria elevam a aplicação para um patamar profissional."
    )
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- REFERÊNCIAS ---
    doc.add_heading('REFERÊNCIAS', level=1)
    doc.add_paragraph("MARTIN, Robert C. Clean Architecture: A Craftsman's Guide to Software Structure and Design. Prentice Hall, 2017.", style='List Bullet')
    doc.add_paragraph("EVANS, Eric. Domain-Driven Design: Tackling Complexity in the Heart of Software. Addison-Wesley, 2003.", style='List Bullet')
    doc.add_paragraph("FOWLER, Martin. Patterns of Enterprise Application Architecture. Addison-Wesley Professional, 2002.", style='List Bullet')
    doc.add_paragraph("GAMMA, Erich et al. Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley, 1994.", style='List Bullet')
    doc.add_paragraph("NEWMAN, Sam. Building Microservices: Designing Fine-Grained Systems. O'Reilly Media, 2015.", style='List Bullet')
    doc.add_paragraph("NESTJS Documentation. Progressive Node.js framework. Disponível em: https://docs.nestjs.com. Acesso em: 14 mai. 2026.", style='List Bullet')
    doc.add_paragraph("PRISMA Documentation. Next-generation ORM for Node.js and TypeScript. Disponível em: https://www.prisma.io/docs. Acesso em: 14 mai. 2026.", style='List Bullet')
    doc.add_paragraph("REDIS Documentation. In-memory data structure store. Disponível em: https://redis.io/docs. Acesso em: 14 mai. 2026.", style='List Bullet')
    doc.add_paragraph("POSTGRESQL Global Development Group. PostgreSQL Documentation. Disponível em: https://www.postgresql.org/docs/. Acesso em: 14 mai. 2026.", style='List Bullet')

    # Salvar o documento
    file_path = "C:/Users/Lucas Sol/Desktop/Relatorio_Tecnico_TechService_OS.docx"
    doc.save(file_path)
    print(f"Relatório gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_report()
